from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_core.retrievers import BaseRetriever
from langchain_community.llms import OpenAI
import pandas as pd
import PyPDF2
import docx
from typing import List
import streamlit as st

def extract_text_from_file(file):
    if file.type == "text/csv":
        df = pd.read_csv(file)
        return df.to_string()
    elif file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif file.type == "text/plain":
        return file.read().decode("utf-8")
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

def split_text_into_chunks(text: str) -> List[str]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

def store_chunks(chunks: List[str], knowledge_base: List[str]):
    knowledge_base.extend(chunks)

class SimpleRetriever(BaseRetriever):
    def __init__(self, knowledge_base: List[str]):
        self.knowledge_base = knowledge_base

    def get_relevant_documents(self, query: str) -> List[Document]:
        return [Document(page_content=doc) for doc in self.knowledge_base if query in doc]

def retrieve_relevant_chunks(query: str, knowledge_base: List[str]) -> List[str]:
    retriever = SimpleRetriever(knowledge_base)
    relevant_docs = retriever.get_relevant_documents(query)
    return [doc.page_content for doc in relevant_docs]

def generate_response_with_knowledge_base(prompt: str, client, model_choice, sub_model_choice, knowledge_base: List[str]) -> str:
    relevant_chunks = retrieve_relevant_chunks(prompt, knowledge_base)
    context = " ".join(relevant_chunks)
    
    if model_choice == "OpenAI":
        response = client.chat.completions.create(
            model=sub_model_choice,
            messages=[
                {"role": "system", "content": context},
                {"role": "user", "content": prompt}
            ],
            stream=False,
        )
        response_content = response['choices'][0]['message']['content']
    elif model_choice == "DeepSeek":
        response = client.generate_response(
            model=sub_model_choice,
            messages=[
                {"role": "system", "content": context},
                {"role": "user", "content": prompt}
            ],
            stream=False,
        )
        response_content = response
    
    return response_content

# Initialize session state for knowledge base
if 'knowledge_base' not in st.session_state:
    st.session_state.knowledge_base = []

uploaded_file = st.file_uploader("Upload a file", type=["csv", "pdf", "txt", "docx"])

if uploaded_file is not None:
    text = extract_text_from_file(uploaded_file)
    chunks = split_text_into_chunks(text)
    store_chunks(chunks, st.session_state.knowledge_base)
    st.success("File processed and chunks stored in knowledge base.")

if st.session_state.knowledge_base:
    if prompt := st.chat_input("What is up?", key="chat_input_key"):
        response = generate_response_with_knowledge_base(prompt, client, model_choice, sub_model_choice, st.session_state.knowledge_base)
        st.write(response)
